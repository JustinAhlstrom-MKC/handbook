#!/usr/bin/env python3
"""
MKC Restaurants Employee Handbook Build Script

Assembles individual policy markdown files into a complete Word document.
Requires: python-docx, pyyaml

Install dependencies:
    pip install python-docx pyyaml

Usage:
    python build/build.py [--exclude-draft] [--output-name NAME]
"""

import os
import re
import yaml
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    exit(1)


def parse_frontmatter(content: str) -> tuple[dict, str]:
    """Extract YAML frontmatter and body content from markdown file."""
    if content.startswith('---'):
        parts = content.split('---', 2)
        if len(parts) >= 3:
            frontmatter = yaml.safe_load(parts[1])
            body = parts[2].strip()
            return frontmatter, body
    return {}, content


def markdown_to_docx_content(doc: Document, content: str, base_heading_level: int = 2):
    """
    Convert markdown content to Word document paragraphs.
    Handles headings, lists, bold, and basic formatting.
    """
    lines = content.split('\n')
    in_list = False

    for line in lines:
        stripped = line.strip()

        # Skip HTML comments (placeholders)
        if stripped.startswith('<!--') and stripped.endswith('-->'):
            continue
        if stripped.startswith('<!--'):
            continue
        if stripped.endswith('-->'):
            continue

        # Empty line
        if not stripped:
            if in_list:
                in_list = False
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # Adjust heading level (H1 in policy becomes H2 in doc, etc.)
            doc_level = min(level + base_heading_level - 1, 9)
            p = doc.add_heading(text, level=doc_level)
            continue

        # Unordered list items
        list_match = re.match(r'^[-*]\s+(.+)$', stripped)
        if list_match:
            in_list = True
            text = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue

        # Horizontal rule (section break)
        if stripped in ['---', '***', '___']:
            doc.add_paragraph()
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)


def add_formatted_text(paragraph, text: str):
    """Add text with basic markdown formatting (bold, italic) to a paragraph."""
    # Pattern to match **bold** and *italic* text
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        # Add formatted text
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def create_cover_page(doc: Document, config: dict):
    """Create a cover page for the handbook."""
    # Add some spacing
    for _ in range(6):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config['company']['name'])
    run.bold = True
    run.font.size = Pt(36)

    doc.add_paragraph()

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config['handbook']['title'])
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()
    doc.add_paragraph()

    # Version and date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version {config['handbook']['version']}")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Effective: {config['handbook']['effective_date']}")
    run.font.size = Pt(14)

    # Page break after cover
    doc.add_page_break()


def create_toc_placeholder(doc: Document):
    """Add a table of contents placeholder."""
    doc.add_heading('Table of Contents', level=1)
    p = doc.add_paragraph()
    p.add_run('[Table of Contents - Update field after opening in Word]')
    p.italic = True
    doc.add_page_break()


def build_handbook(config_path: str, exclude_draft: bool = False, output_name: str = None):
    """Build the complete handbook document."""

    # Load configuration
    config_path = Path(config_path)
    base_dir = config_path.parent

    with open(config_path, 'r') as f:
        config = yaml.safe_load(f)

    # Create document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Create cover page
    if config['build'].get('include_cover', True):
        create_cover_page(doc, config)

    # Create TOC placeholder
    if config['build'].get('include_toc', True):
        create_toc_placeholder(doc)

    # Process each section
    policies_dir = base_dir / 'policies'

    for section in config['sections']:
        section_id = section['id']
        section_title = section['title']
        section_path = policies_dir / section_id

        # Add section heading
        doc.add_heading(section_title, level=1)

        # Process each policy in the section
        for policy_file in section['policies']:
            policy_path = section_path / policy_file

            if not policy_path.exists():
                print(f"Warning: Policy file not found: {policy_path}")
                continue

            with open(policy_path, 'r') as f:
                content = f.read()

            frontmatter, body = parse_frontmatter(content)

            # Skip draft policies if requested
            if exclude_draft and frontmatter.get('status') == 'draft':
                print(f"Skipping draft policy: {policy_file}")
                continue

            # Add policy content
            markdown_to_docx_content(doc, body)

            # Add some spacing between policies
            doc.add_paragraph()

        # Page break after each section (except the last)
        if section != config['sections'][-1]:
            doc.add_page_break()

    # Save document
    output_dir = base_dir / config['build']['output_dir']
    output_dir.mkdir(exist_ok=True)

    if output_name:
        output_filename = f"{output_name}.docx"
    else:
        date_str = datetime.now().strftime('%Y-%m-%d')
        version = config['handbook']['version']
        output_filename = f"MKC-Employee-Handbook-v{version}-{date_str}.docx"

    output_path = output_dir / output_filename
    doc.save(output_path)

    print(f"Handbook built successfully: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description='Build MKC Employee Handbook')
    parser.add_argument('--exclude-draft', action='store_true',
                        help='Exclude policies with status: draft')
    parser.add_argument('--output-name', type=str,
                        help='Custom output filename (without extension)')
    parser.add_argument('--config', type=str, default='config.yaml',
                        help='Path to config file (default: config.yaml)')

    args = parser.parse_args()

    # Find config file
    script_dir = Path(__file__).parent.parent
    config_path = script_dir / args.config

    if not config_path.exists():
        print(f"Error: Config file not found: {config_path}")
        exit(1)

    build_handbook(config_path, args.exclude_draft, args.output_name)


if __name__ == '__main__':
    main()
